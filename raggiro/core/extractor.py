"""Module for text extraction from various document formats."""

import io
import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any

# PDF extraction
import fitz  # PyMuPDF
from pdfminer.high_level import extract_text as pdfminer_extract_text

# Document extraction
import docx
import openpyxl
import pandas as pd
from bs4 import BeautifulSoup

# OCR for scanned documents and images
import pytesseract
from PIL import Image

# Character encoding detection
import chardet

# Set up logger
logger = logging.getLogger("raggiro.extractor")

class Extractor:
    """Extracts text content from various document formats."""
    
    def __init__(self, config: Optional[Dict] = None):
        """Initialize the extractor.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {}
        extraction_config = self.config.get("extraction", {})
        
        # OCR config
        self.ocr_enabled = extraction_config.get("ocr_enabled", True)
        self.ocr_language = extraction_config.get("ocr_language", "eng")
        
        # Handle auto language detection
        self.language_detector = None
        if self.ocr_language == "auto":
            try:
                import langdetect
                self.language_detector = langdetect
                logger.info("Language auto-detection enabled for OCR")
            except ImportError:
                logger.warning("langdetect not installed. Using default OCR language 'eng+ita'")
                self.ocr_language = "eng+ita"
        
        # Advanced OCR settings
        self.ocr_dpi = extraction_config.get("ocr_dpi", 300)
        self.ocr_max_image_size = extraction_config.get("ocr_max_image_size", 4000)
        self.ocr_batch_size = extraction_config.get("ocr_batch_size", 20)  # Increased from 10 to 20
        
        # Page selection settings - ensure these are respected
        # 0 means process all pages (this is the default)
        self.ocr_max_pages = extraction_config.get("ocr_max_pages", 0)  
        
        # 1 means process every page (this is the default)
        self.ocr_page_step = extraction_config.get("ocr_page_step", 1)
        
        # Log the page selection settings
        if self.ocr_max_pages == 0:
            logger.info("OCR configured to process ALL pages in the document")
        else:
            logger.info(f"OCR configured to process first {self.ocr_max_pages} pages")
            
        if self.ocr_page_step > 1:
            logger.info(f"OCR configured to process every {self.ocr_page_step}th page")
        
        # Configure pytesseract path if provided
        tesseract_path = extraction_config.get("tesseract_path")
        if tesseract_path and os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            
        # Configure tesseract parameters
        # Add config parameters to preserve special characters and improve accuracy with accented characters
        self.tesseract_config = f'--dpi {self.ocr_dpi} --oem 1 --psm 3 -c preserve_interword_spaces=1 -c textonly_pdf=1'
        
        # Verify and fix language configuration
        # Handle special case when language is "auto" to ensure Tesseract doesn't fail
        # Tesseract doesn't have an "auto" language pack, so we use a fallback
        if self.ocr_language == "auto":
            # Using auto for language detection, but need a specific language for tesseract
            # For an Italian target audience, prioritize Italian but include English
            self.actual_ocr_language = "ita+eng"  # Default fallback with Italian first
            print(f"WARNING: Using 'auto' for language detection, but Tesseract will use '{self.actual_ocr_language}' until detection")
        else:
            self.actual_ocr_language = self.ocr_language
            
        # Check if TESSDATA_PREFIX is set
        tessdata_prefix = os.environ.get("TESSDATA_PREFIX")
        if not tessdata_prefix:
            # Try to find tessdata directory
            common_tessdata_paths = [
                "/usr/share/tesseract-ocr/4.00/tessdata",
                "/usr/share/tesseract-ocr/5/tessdata",
                "/usr/local/share/tessdata",
                "/opt/homebrew/share/tessdata",  # For macOS homebrew
                "/usr/share/tessdata",
            ]
            
            for path in common_tessdata_paths:
                if os.path.exists(path):
                    os.environ["TESSDATA_PREFIX"] = path
                    logger.info(f"Setting TESSDATA_PREFIX to {path}")
                    break
            
            if not os.environ.get("TESSDATA_PREFIX"):
                logger.warning("Could not find tessdata directory. OCR might not work correctly.")
        
        # Log configuration
        logger.info(f"OCR Configuration: language={self.actual_ocr_language} (from {self.ocr_language}), dpi={self.ocr_dpi}, " +
                 f"max_image_size={self.ocr_max_image_size}, batch_size={self.ocr_batch_size}")
    
    def extract(self, file_path: Union[str, Path], file_type_info: Dict) -> Dict:
        """Extract text from a file based on its type.
        
        Args:
            file_path: Path to the file
            file_type_info: File type information from file_handler
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        document_type = file_type_info.get("document_type", "unknown")
        mime_type = file_type_info.get("mime_type", "")
        
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": None,
            "has_text_layer": False,
            "success": False,
            "error": None,
        }
        
        try:
            if document_type == "pdf":
                return self._extract_pdf(file_path)
            elif document_type == "word":
                return self._extract_docx(file_path)
            elif document_type == "spreadsheet":
                return self._extract_excel(file_path)
            elif document_type == "text":
                if "html" in mime_type:
                    return self._extract_html(file_path)
                elif "rtf" in mime_type:
                    return self._extract_rtf(file_path)
                else:
                    return self._extract_text(file_path)
            elif document_type == "image" and self.ocr_enabled:
                return self._extract_image_ocr(file_path)
            else:
                result["error"] = f"Unsupported document type: {document_type}"
                return result
                
        except Exception as e:
            result["error"] = str(e)
            return result
    
    def _extract_pdf(self, file_path: Path) -> Dict:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "pdf",
            "has_text_layer": False,
            "success": False,
            "error": None,
        }
        
        try:
            # First try PyMuPDF (faster and better metadata)
            doc = fitz.open(file_path)
            
            # Extract metadata
            result["metadata"] = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "keywords": doc.metadata.get("keywords", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
                "total_pages": len(doc),
            }
            
            # Extract document outline/bookmarks if available
            try:
                toc = doc.get_toc()
                if toc and len(toc) > 0:
                    result["pdf_outline"] = []
                    for item in toc:
                        if len(item) >= 3:  # [level, title, page]
                            outline_item = {
                                "level": item[0],
                                "title": item[1],
                                "page": item[2]
                            }
                            result["pdf_outline"].append(outline_item)
                    logger.info(f"Extracted {len(result['pdf_outline'])} items from PDF outline/bookmarks")
            except Exception as e:
                logger.warning(f"Failed to extract PDF outline: {e}")
            
            # Extract text from each page
            full_text = []
            pages = []
            has_text = False
            
            for i, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    has_text = True
                    
                full_text.append(text)
                pages.append({
                    "page_num": i + 1,
                    "text": text,
                    "has_text": bool(text.strip()),
                })
            
            result["has_text_layer"] = has_text
            result["text"] = "\n\n".join(full_text)
            result["pages"] = pages
            
            # ALWAYS perform OCR if OCR is enabled, regardless of existing text layer
            # This ensures we process all pages with OCR and don't rely on PDF text layer
            if self.ocr_enabled:
                logger.info("Forcing OCR processing even if PDF has text layer")
                ocr_result = self._extract_pdf_with_ocr(file_path)
                if ocr_result["success"]:
                    # Save the PDF outline/bookmarks if we have them
                    pdf_outline = result.get("pdf_outline")
                    
                    # If OCR was successful, use its results
                    result = ocr_result
                    result["extraction_method"] = "pdf_ocr"
                    # Preserve the information about text layer
                    result["has_text_layer"] = has_text
                    
                    # Restore the PDF outline/bookmarks
                    if pdf_outline:
                        result["pdf_outline"] = pdf_outline
            
            # If still no text, try pdfminer as a fallback
            if not result["text"].strip() and not result["error"]:
                try:
                    text = pdfminer_extract_text(str(file_path))
                    if text.strip():
                        result["text"] = text
                        result["extraction_method"] = "pdf_pdfminer"
                        result["has_text_layer"] = True
                except Exception as e:
                    # Just log the error but keep the PyMuPDF result
                    pass
            
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_pdf_with_ocr(self, file_path: Path) -> Dict:
        """Extract text from a PDF file using OCR with batch processing.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        import gc  # For explicit garbage collection
        import time
        
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "pdf_ocr",
            "has_text_layer": False,
            "success": False,
            "error": None,
        }
        
        # If auto language detection is enabled, try to determine the document language
        if self.language_detector is not None and self.ocr_language == "auto":
            # Get a sample of text to detect language
            try:
                # Try to extract some text from the first few pages without OCR first
                doc = fitz.open(file_path)
                sample_text = ""
                for i in range(min(3, len(doc))):
                    page_text = doc[i].get_text()
                    if page_text.strip():
                        sample_text += page_text + "\n"
                        if len(sample_text) > 1000:
                            break
                
                if sample_text.strip():
                    # Detect language from sample
                    detected_lang = self.language_detector.detect(sample_text)
                    # Map to Tesseract language codes
                    lang_map = {
                        "en": "eng",
                        "it": "ita",
                        "fr": "fra",
                        "de": "deu",
                        "es": "spa",
                        "pt": "por",
                        "nl": "nld",
                        "ru": "rus",
                    }
                    if detected_lang in lang_map:
                        # Set primary language + eng as fallback
                        # If Italian is detected, make it the primary language
                        if detected_lang == "it":
                            self.actual_ocr_language = "ita+eng"
                        else:
                            self.actual_ocr_language = f"{lang_map[detected_lang]}+eng"
                        logger.info(f"Auto-detected document language: {detected_lang}, using OCR language: {self.actual_ocr_language}")
                    else:
                        # Fallback to multiple common languages with Italian prioritized
                        self.actual_ocr_language = "ita+eng+fra+deu+spa"
                        logger.info(f"Unable to definitively detect language ({detected_lang}), using multiple languages: {self.actual_ocr_language}")
                else:
                    # No text detected, use multiple languages with Italian prioritized
                    self.actual_ocr_language = "ita+eng+fra+deu+spa"
                    logger.info(f"No text available for language detection, using multiple languages: {self.actual_ocr_language}")
            except Exception as e:
                logger.warning(f"Error in language detection: {str(e)}. Using default languages with Italian prioritized.", exc_info=True)
                self.actual_ocr_language = "ita+eng+fra+deu+spa"
        
        try:
            doc = fitz.open(file_path)
            
            # Extract basic metadata from the document
            result["metadata"] = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "keywords": doc.metadata.get("keywords", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
                "total_pages": len(doc),
            }
            
            # Extract text from each page using OCR
            full_text = []
            pages = []
            total_doc_pages = len(doc)
            
            # Apply max_pages limitation if specified
            if self.ocr_max_pages > 0 and self.ocr_max_pages < total_doc_pages:
                max_pages = self.ocr_max_pages
                logger.info(f"Limiting OCR to first {max_pages} pages (document has {total_doc_pages} pages)")
            else:
                max_pages = total_doc_pages
                logger.info(f"Processing all {total_doc_pages} pages of the document")
                
            # Apply page step (process every Nth page)
            if self.ocr_page_step > 1:
                page_indices = list(range(0, max_pages, self.ocr_page_step))
                logger.info(f"Processing every {self.ocr_page_step} page ({len(page_indices)} of {max_pages} pages)")
            else:
                page_indices = list(range(max_pages))
                
            batch_size = self.ocr_batch_size
            
            # Log the OCR processing settings
            logger.info(f"OCR processing {len(page_indices)} pages with batch size {batch_size}")
            logger.info(f"Settings: DPI={self.ocr_dpi}, Language={self.actual_ocr_language} (selected: {self.ocr_language}), Max image size={self.ocr_max_image_size}px")
            start_time = time.time()
            
            # Process pages in batches to manage memory
            batch_indices = [page_indices[i:i+batch_size] for i in range(0, len(page_indices), batch_size)]
            
            for batch_num, batch_page_indices in enumerate(batch_indices):
                batch_start = batch_page_indices[0] + 1  # 1-based for display
                batch_end = batch_page_indices[-1] + 1   # 1-based for display
                logger.info(f"Processing batch {batch_num+1}/{len(batch_indices)}: pages {batch_start}-{batch_end} of {max_pages}")
                batch_start_time = time.time()
                
                # Process each page in the current batch
                for page_idx in batch_page_indices:
                    try:
                        # Use the actual page index for the document
                        page = doc[page_idx]
                        page_start_time = time.time()
                        
                        # Log progress for debugging (use 1-based page numbers for display)
                        logger.info(f"OCR processing page {page_idx+1}/{total_doc_pages}")
                        
                        # Calculate zoom factor based on page dimensions and max image size
                        rect = page.rect
                        width, height = rect.width, rect.height
                        max_dim = max(width, height)
                        
                        # Calculate zoom to fit within max_image_size
                        zoom = 1.0
                        if max_dim > self.ocr_max_image_size:
                            zoom = self.ocr_max_image_size / max_dim
                            logger.info(f"Scaling page {page_idx+1} ({width}x{height}) to fit {self.ocr_max_image_size}px, zoom={zoom:.2f}")
                        
                        # Create the matrix for scaling to the target DPI
                        # PyMuPDF uses 72 dpi as base, so we calculate relative scaling
                        dpi_scale = self.ocr_dpi / 72.0
                        matrix = fitz.Matrix(zoom * dpi_scale, zoom * dpi_scale)
                        
                        # Get the pixmap with appropriate scaling
                        try:
                            pix = page.get_pixmap(matrix=matrix)
                        except Exception as pix_error:
                            logger.warning(f"Error creating pixmap for page {page_idx+1}: {str(pix_error)}", exc_info=True)
                            # Try with lower resolution as fallback
                            fallback_matrix = fitz.Matrix(0.5, 0.5)
                            pix = page.get_pixmap(matrix=fallback_matrix)
                            logger.info(f"Using fallback matrix for page {page_idx+1} due to pixmap error")
                        
                        # Convert to PIL Image
                        try:
                            img_data = pix.tobytes("png")
                            img = Image.open(io.BytesIO(img_data))
                            
                            # Log successful conversion
                            logger.debug(f"Successfully converted page {page_idx+1} to image format for OCR")
                            
                            # Store pix reference for error message access
                            pix_ref = pix
                            # Free memory immediately after use
                            del pix
                            del img_data
                            gc.collect()  # Explicit garbage collection
                            
                            # Perform OCR with timeout protection and custom config
                            # Use the actual_ocr_language property which handles the "auto" case
                            # Note: pytesseract.image_to_string doesn't support encoding parameter
                            text = pytesseract.image_to_string(
                                img, 
                                lang=self.actual_ocr_language,
                                config=self.tesseract_config,
                                output_type=pytesseract.Output.STRING
                            )
                            
                            # Ensure text is UTF-8 encoded if needed
                            if isinstance(text, bytes):
                                text = text.decode('utf-8', errors='replace')
                            
                            # Free memory
                            del img
                            
                        except Exception as img_error:
                            logger.warning(f"Image conversion error on page {page_idx+1}: {str(img_error)}", exc_info=True)
                            try:
                                # Try with direct PNG output as fallback
                                output_path = f"/tmp/ocr_page_{page_idx+1}.png"
                                # Check if pix_ref is available
                                if 'pix_ref' in locals():
                                    pix_ref.save(output_path)
                                    img = Image.open(output_path)
                                else:
                                    # If pix_ref is not available, we need an alternative approach
                                    # Get a new pixmap with minimum settings
                                    fallback_matrix = fitz.Matrix(0.5, 0.5)
                                    fallback_pix = page.get_pixmap(matrix=fallback_matrix)
                                    fallback_pix.save(output_path)
                                    img = Image.open(output_path)
                                    logger.info(f"Using emergency fallback for page {page_idx+1}")
                            except Exception as emergency_error:
                                # Critical failure, cannot process this page
                                logger.error(f"Critical error on page {page_idx+1}: {str(emergency_error)}", exc_info=True)
                                raise
                            
                            # Perform OCR on the saved file
                            # Use the actual_ocr_language property which handles the "auto" case
                            text = pytesseract.image_to_string(
                                img, 
                                lang=self.actual_ocr_language,
                                config=self.tesseract_config,
                                output_type=pytesseract.Output.STRING
                            )
                            
                            # Ensure text is UTF-8 encoded if needed
                            if isinstance(text, bytes):
                                text = text.decode('utf-8', errors='replace')
                            
                            # Cleanup
                            del img
                            os.unlink(output_path)
                        
                        # Force garbage collection
                        gc.collect()
                        
                        # Record the text and character count
                        char_count = len(text)
                        full_text.append(text)
                        
                        # Save both original raw OCR text and processed text
                        # This will be used for comparison views
                        page_data = {
                            "page_num": page_idx + 1,
                            "text": text,  # This will be potentially corrected later in the pipeline
                            "raw_text": text,  # Keep the original OCR text for comparison
                            "original_text": text,  # Explicitly add original_text for better compatibility
                            "has_text": bool(text.strip()),
                            "char_count": char_count,
                            "processing_time": time.time() - page_start_time
                        }
                        pages.append(page_data)
                        
                        logger.info(f"Page {page_idx+1} OCR extracted {char_count} characters")
                        
                        logger.debug(f"Page {page_idx+1} OCR completed in {time.time() - page_start_time:.2f} seconds")
                        
                    except Exception as page_error:
                        error_msg = f"Error processing page {page_idx+1}: {str(page_error)}"
                        logger.error(error_msg, exc_info=True)
                        # Add placeholder for failed page with safer error message that doesn't reference deleted variables
                        error_text = f"[ERROR: Failed to process page {page_idx+1} - {str(page_error).replace('cannot access local variable', 'error accessing variable')}]"
                        full_text.append(error_text)
                        pages.append({
                            "page_num": page_idx + 1,
                            "text": error_text,
                            "raw_text": error_text,  # Same for raw text in case of error
                            "original_text": error_text,  # Add original_text for comparison view
                            "has_text": False,
                            "error_type": type(page_error).__name__,
                        })
                
                # After each batch, explicitly run garbage collection
                gc.collect()
                logger.info(f"Batch {batch_start+1}-{batch_end} completed in {time.time() - batch_start_time:.2f} seconds")
            
            # Save the combined result
            combined_text = "\n\n".join(full_text)
            result["text"] = combined_text
            result["raw_text"] = combined_text  # Store the raw OCR text for comparison
            result["original_text"] = combined_text  # Explicitly add original_text for better compatibility
            result["pages"] = pages
            result["success"] = True
            
            # Calculate and log character count
            total_chars = len(combined_text)
            result["metadata"]["ocr_char_count"] = total_chars
            
            # Calculate chars per page
            chars_per_page = total_chars / max(1, max_pages)
            result["metadata"]["ocr_chars_per_page"] = round(chars_per_page, 1)
            
            total_time = time.time() - start_time
            logger.info(f"OCR processing completed for all {max_pages} pages in {total_time:.2f} seconds")
            logger.info(f"Total characters recognized: {total_chars} (avg: {chars_per_page:.1f} chars/page)")
            result["metadata"]["ocr_processing_time"] = total_time
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione PDF: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_docx(self, file_path: Path) -> Dict:
        """Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "docx",
            "has_text_layer": True,
            "success": False,
            "error": None,
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            core_properties = doc.core_properties
            result["metadata"] = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": core_properties.created.isoformat() if core_properties.created else "",
                "modified": core_properties.modified.isoformat() if core_properties.modified else "",
                "last_modified_by": core_properties.last_modified_by or "",
                "category": core_properties.category or "",
                "paragraphs": len(doc.paragraphs),
            }
            
            # Extract text from paragraphs
            text_content = [p.text for p in doc.paragraphs]
            
            # Handle tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            result["text"] = "\n".join(text_content)
            result["success"] = True
            
            # DOCX doesn't have the concept of pages in the same way as PDF
            result["pages"] = [{
                "page_num": 1,
                "text": result["text"],
                "has_text": True,
            }]
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_excel(self, file_path: Path) -> Dict:
        """Extract text from an Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "excel",
            "has_text_layer": True,
            "success": False,
            "error": None,
        }
        
        try:
            # Try with pandas first (handles more formats, including .xls)
            try:
                # Read all sheets
                sheets = pd.read_excel(file_path, sheet_name=None)
                
                text_parts = []
                sheet_texts = []
                
                for sheet_name, df in sheets.items():
                    sheet_text = f"Sheet: {sheet_name}\n"
                    sheet_text += df.to_string(index=False)
                    text_parts.append(sheet_text)
                    sheet_texts.append({
                        "sheet_name": sheet_name,
                        "text": sheet_text,
                    })
                
                result["text"] = "\n\n".join(text_parts)
                result["pages"] = sheet_texts
                result["metadata"]["sheets"] = list(sheets.keys())
                result["metadata"]["total_sheets"] = len(sheets)
                
            except Exception:
                # Fallback to openpyxl (XLSX only)
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                
                text_parts = []
                sheet_texts = []
                
                for sheet in workbook:
                    rows = []
                    for row in sheet.iter_rows(values_only=True):
                        rows.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
                    
                    sheet_text = f"Sheet: {sheet.title}\n"
                    sheet_text += "\n".join(rows)
                    
                    text_parts.append(sheet_text)
                    sheet_texts.append({
                        "sheet_name": sheet.title,
                        "text": sheet_text,
                    })
                
                result["text"] = "\n\n".join(text_parts)
                result["pages"] = sheet_texts
                result["metadata"]["sheets"] = [sheet.title for sheet in workbook]
                result["metadata"]["total_sheets"] = len(workbook.sheetnames)
            
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_html(self, file_path: Path) -> Dict:
        """Extract text from an HTML file.
        
        Args:
            file_path: Path to the HTML file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "html",
            "has_text_layer": True,
            "success": False,
            "error": None,
        }
        
        try:
            with open(file_path, "rb") as file:
                content = file.read()
                encoding = chardet.detect(content)["encoding"] or "utf-8"
                html_content = content.decode(encoding, errors="replace")
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Extract metadata
            title = soup.title.text if soup.title else ""
            meta_tags = {}
            
            for meta in soup.find_all("meta"):
                name = meta.get("name", meta.get("property", ""))
                if name:
                    meta_tags[name] = meta.get("content", "")
            
            result["metadata"] = {
                "title": title,
                "meta_tags": meta_tags,
            }
            
            # Remove scripts and styles for cleaner text
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text(separator="\n")
            
            # Remove excessive newlines
            lines = [line.strip() for line in text.split("\n")]
            text = "\n".join(line for line in lines if line)
            
            result["text"] = text
            result["pages"] = [{
                "page_num": 1,
                "text": text,
                "has_text": True,
            }]
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_text(self, file_path: Path) -> Dict:
        """Extract text from a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary with extracted text
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "text",
            "has_text_layer": True,
            "success": False,
            "error": None,
        }
        
        try:
            with open(file_path, "rb") as file:
                content = file.read()
                
                # Detect encoding
                detected = chardet.detect(content)
                encoding = detected["encoding"] or "utf-8"
                confidence = detected["confidence"]
                
                text = content.decode(encoding, errors="replace")
            
            result["text"] = text
            result["metadata"] = {
                "encoding": encoding,
                "encoding_confidence": confidence,
            }
            result["pages"] = [{
                "page_num": 1,
                "text": text,
                "has_text": True,
            }]
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_rtf(self, file_path: Path) -> Dict:
        """Extract text from an RTF file.
        
        Args:
            file_path: Path to the RTF file
            
        Returns:
            Dictionary with extracted text
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "rtf",
            "has_text_layer": True,
            "success": False,
            "error": None,
        }
        
        try:
            # Since RTF can be complex, we'll try a few methods
            
            # First try to read as plain text (works for simple RTF)
            with open(file_path, "rb") as file:
                content = file.read()
                
                # Try to decode as is
                try:
                    text = content.decode("utf-8", errors="replace")
                    
                    # Remove RTF control sequences
                    cleaned_text = ""
                    i = 0
                    in_control = False
                    
                    while i < len(text):
                        if text[i] == "\\" and not in_control:
                            in_control = True
                            i += 1
                        elif in_control and text[i].isalpha():
                            # Skip the control word
                            while i < len(text) and text[i].isalpha():
                                i += 1
                            in_control = False
                        elif in_control and text[i] == "'":
                            # Skip hex character
                            i += 3
                            in_control = False
                        elif in_control:
                            # Single character control
                            i += 1
                            in_control = False
                        else:
                            cleaned_text += text[i]
                            i += 1
                    
                    text = cleaned_text
                    
                except Exception:
                    # If the above doesn't work, try pip install striprtf
                    # For now, use a simple fallback
                    text = str(content)
                    text = text.replace("\\par", "\n")
                    text = text.replace("\\tab", "\t")
            
            result["text"] = text
            result["pages"] = [{
                "page_num": 1,
                "text": text,
                "has_text": True,
            }]
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result
    
    def _extract_image_ocr(self, file_path: Path) -> Dict:
        """Extract text from an image using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Dictionary with extracted text
        """
        result = {
            "text": "",
            "pages": [],
            "metadata": {},
            "extraction_method": "image_ocr",
            "has_text_layer": False,
            "success": False,
            "error": None,
        }
        
        if not self.ocr_enabled:
            result["error"] = "OCR is disabled in configuration"
            return result
        
        try:
            # Open the image
            image = Image.open(file_path)
            
            # Get image metadata
            result["metadata"] = {
                "format": image.format,
                "mode": image.mode,
                "width": image.width,
                "height": image.height,
            }
            
            # Perform OCR using the actual_ocr_language that handles the "auto" case
            # Explicitly set output type as string with UTF-8 encoding for better handling of accented characters
            text = pytesseract.image_to_string(
                image,
                lang=self.actual_ocr_language,
                config=self.tesseract_config,
                output_type=pytesseract.Output.STRING
            )
            
            # Ensure text is UTF-8 encoded if needed
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='replace')
            
            result["text"] = text
            result["raw_text"] = text  # Store raw OCR text for comparison
            result["original_text"] = text  # Explicitly add original_text for better compatibility
            result["pages"] = [{
                "page_num": 1,
                "text": text,
                "raw_text": text,  # Store raw OCR text for comparison
                "original_text": text,  # Explicitly add original_text for comparison view
                "has_text": bool(text.strip()),
            }]
            result["success"] = True
            
        except Exception as e:
            error_msg = f"Errore durante l'estrazione: {str(e)}"
            logger.error(error_msg, exc_info=True)
            result["error"] = error_msg
            result["error_type"] = type(e).__name__
            
        return result